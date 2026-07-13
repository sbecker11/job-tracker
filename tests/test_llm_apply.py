"""Tests for LLM-driven JD evaluation + résumé/cover-letter generation
(pipeline/llm_apply.py).

No real Anthropic API calls are made here — the client is always a fake
stand-in so the test suite runs offline and free of charge. The candidate
profile is also faked via an autouse fixture so tests don't depend on
~/CLAUDE.md existing on the machine running them.
"""

from __future__ import annotations

import json

import pytest
from docx import Document

from job_tracker.pipeline import llm_apply


class _FakeBlock:
    def __init__(self, text: str):
        self.type = "text"
        self.text = text


class _FakeUsage:
    def __init__(self, input_tokens: int, output_tokens: int):
        self.input_tokens = input_tokens
        self.output_tokens = output_tokens


class _FakeMessage:
    def __init__(self, text: str, input_tokens: int, output_tokens: int):
        self.content = [_FakeBlock(text)]
        self.usage = _FakeUsage(input_tokens, output_tokens)


class _FakeClient:
    """Stands in for anthropic.Anthropic. Scripted with a queue of responses
    consumed in call order, so multi-call flows (JSON repair, house-rule
    repair, evaluate-then-generate) can be tested deterministically."""

    def __init__(self, responses: list | None = None, error: Exception | None = None):
        # Each item is either a raw text string (defaults: 100 in / 50 out
        # tokens) or an (text, input_tokens, output_tokens) tuple.
        self.responses = list(responses or [])
        self.error = error
        self.calls: list[dict] = []
        self.messages = self

    def create(self, **kwargs):
        self.calls.append(kwargs)
        if self.error is not None:
            raise self.error
        if not self.responses:
            raise AssertionError("FakeClient ran out of scripted responses")
        item = self.responses.pop(0)
        if isinstance(item, tuple):
            text, input_tokens, output_tokens = item
        else:
            text, input_tokens, output_tokens = item, 100, 50
        return _FakeMessage(text, input_tokens, output_tokens)


@pytest.fixture(autouse=True)
def fake_profile(monkeypatch, tmp_path):
    """Point the module at a small throwaway candidate profile for every
    test, so nothing here depends on ~/CLAUDE.md existing."""
    profile_path = tmp_path / "CLAUDE.md"
    profile_path.write_text("FAKE CANDIDATE PROFILE FOR TESTS", encoding="utf-8")
    monkeypatch.setattr(llm_apply, "_CANDIDATE_PROFILE_PATH", profile_path)
    return profile_path


_EVAL_PAYLOAD_PURSUE = json.dumps(
    {
        "job_summary": "A backend role building agentic AI systems.",
        "dealbreaker_checks": [{"check": "Banned stack", "status": "clean", "notes": "Python/TypeScript only."}],
        "skills_alignment": [{"requirement": "Python", "evidence": "Years of Python work.", "strength": "strong"}],
        "match_pct": 85,
        "flags": ["Title reads senior but scope may be mid-level."],
        "structural_verdict": "PASS on structure",
        "next_step": "",
        "verdict": "pursue",
        "rationale": "Strong overall fit.",
        "framing_guidance": ["Lead with the agentic-AI throughline."],
        "cover_letter_strategy": "Frame the agentic-AI throughline as the core narrative.",
        "interview_prep": ["Open with the agentic-AI project as the go-to example."],
    }
)

_EVAL_PAYLOAD_PASS = json.dumps(
    {
        "job_summary": "A frontend-heavy role on a small team.",
        "dealbreaker_checks": [
            {"check": "Banned stack", "status": "fail", "notes": "Angular required as sole frontend framework."}
        ],
        "skills_alignment": [],
        "match_pct": 10,
        "flags": [],
        "structural_verdict": "FAIL on structure",
        "next_step": "",
        "verdict": "pass",
        "rationale": "Dealbreaker fires.",
        "framing_guidance": [],
        "cover_letter_strategy": "",
        "interview_prep": [],
    }
)

_CLEAN_CONTENT = {
    "resume": {
        "positioning_line": "Senior Software Engineer",
        "summary": "A summary.",
        "skills": ["Python", "React"],
        "experience": [
            {"employer": "HomePortfolio", "dates": "1997-2002", "role_note": "CTO", "bullets": ["Did things."], "subsections": []},
            {"employer": "Spexture (Independent Consulting)", "dates": "2019-present", "role_note": None, "bullets": ["Built things."], "subsections": [
                {"heading": "Portfolio Projects", "bullets": ["Built linkage-engine."]},
            ]},
        ],
        "education": ["PhD, Computer Vision — MIT Media Lab"],
    },
    "cover_letter": {
        "salutation": "Dear Hiring Team,",
        "paragraphs": ["I am excited to apply.", "I have relevant experience."],
    },
}


# --- evaluate_lead ---------------------------------------------------------


def test_evaluate_lead_parses_response_and_computes_metrics():
    client = _FakeClient(responses=[(_EVAL_PAYLOAD_PURSUE, 1000, 200)])
    result = llm_apply.evaluate_lead("some JD text", company="Acme", title="Engineer", model="claude-haiku-4-5", client=client)

    assert result.verdict == "pursue"
    assert result.match_pct == 85
    assert result.job_summary == "A backend role building agentic AI systems."
    assert result.dealbreaker_checks == [{"check": "Banned stack", "status": "clean", "notes": "Python/TypeScript only."}]
    assert result.skills_alignment == [{"requirement": "Python", "evidence": "Years of Python work.", "strength": "strong"}]
    assert result.flags == ["Title reads senior but scope may be mid-level."]
    assert result.rationale == "Strong overall fit."
    assert result.framing_guidance == ["Lead with the agentic-AI throughline."]
    assert result.structural_verdict == "PASS on structure"
    assert result.cover_letter_strategy == "Frame the agentic-AI throughline as the core narrative."
    assert result.interview_prep == ["Open with the agentic-AI project as the go-to example."]
    assert len(client.calls) == 1
    assert client.calls[0]["model"] == "claude-haiku-4-5"

    assert result.metrics is not None
    assert result.metrics.step == "evaluate"
    assert result.metrics.input_tokens == 1000
    assert result.metrics.output_tokens == 200
    assert result.metrics.cost_usd == pytest.approx(1000 / 1_000_000 * 1.00 + 200 / 1_000_000 * 5.00)


def test_evaluate_lead_defaults_verdict_to_review_when_missing():
    payload = json.dumps({"match_pct": 50})
    client = _FakeClient(responses=[payload])
    result = llm_apply.evaluate_lead("JD", company="Acme", title="Eng", model="claude-haiku-4-5", client=client)
    assert result.verdict == "review"


# --- JSON parsing / repair retry --------------------------------------------


def test_call_and_parse_json_repairs_invalid_json_with_one_retry():
    broken = '{"a": "unescaped " quote"}'
    fixed = json.dumps({"a": "fixed"})
    client = _FakeClient(responses=[(broken, 500, 500), (fixed, 200, 50)])

    data, calls = llm_apply._call_and_parse_json(
        "system", "user", model="claude-haiku-4-5", client=client, step="evaluate"
    )

    assert data == {"a": "fixed"}
    assert len(client.calls) == 2
    assert [c.step for c in calls] == ["evaluate", "evaluate_json_repair"]


def test_call_and_parse_json_strips_markdown_fence():
    payload = "```json\n" + json.dumps({"x": 1}) + "\n```"
    client = _FakeClient(responses=[payload])
    data, calls = llm_apply._call_and_parse_json("s", "u", model="claude-haiku-4-5", client=client, step="generate")
    assert data == {"x": 1}
    assert len(calls) == 1


def test_call_and_parse_json_raises_if_repair_also_fails():
    client = _FakeClient(responses=["not json", "still not json"])
    with pytest.raises(json.JSONDecodeError):
        llm_apply._call_and_parse_json("s", "u", model="claude-haiku-4-5", client=client, step="evaluate")


# --- house rule checks -------------------------------------------------------


def test_check_house_rules_flags_banned_term():
    content = {"resume": {"summary": "Worked at Cambria Health for years."}, "cover_letter": {}}
    warnings = llm_apply._check_house_rules(content, company="Acme")
    assert any("Cambria" in w for w in warnings)


def test_check_house_rules_flags_compensation_figure():
    content = {
        "resume": {},
        "cover_letter": {"paragraphs": ["I target $90/hr on W2, roughly $105/hr on C2C."]},
    }
    warnings = llm_apply._check_house_rules(content, company="Bellese")
    assert any("compensation" in w for w in warnings)


def test_check_house_rules_flags_work_authorization_statement():
    content = {
        "resume": {},
        "cover_letter": {"paragraphs": ["I am a US citizen and eligible for Public Trust clearance."]},
    }
    warnings = llm_apply._check_house_rules(content, company="Bellese")
    assert any("work-authorization" in w for w in warnings)


def test_check_house_rules_clean_content_has_no_warnings():
    warnings = llm_apply._check_house_rules(_CLEAN_CONTENT, company="Acme")
    assert warnings == []


# --- house rule auto-repair --------------------------------------------------


def test_repair_house_rule_violations_calls_model_and_returns_fixed_content():
    dirty = {"cover_letter": {"paragraphs": ["I am a US citizen."]}}
    clean = {"cover_letter": {"paragraphs": ["I bring relevant experience."]}}
    client = _FakeClient(responses=[(json.dumps(clean), 300, 100)])

    repaired, calls = llm_apply._repair_house_rule_violations(
        dirty, issues=["possible work-authorization statement"], model="claude-haiku-4-5", client=client
    )

    assert repaired == clean
    assert len(calls) == 1
    assert calls[0].step == "generate_house_rule_repair"


def test_repair_house_rule_violations_keeps_original_on_failure():
    dirty = {"cover_letter": {"paragraphs": ["I am a US citizen."]}}
    client = _FakeClient(responses=["not json", "still not json"])
    repaired, calls = llm_apply._repair_house_rule_violations(
        dirty, issues=["issue"], model="claude-haiku-4-5", client=client
    )
    assert repaired == dirty
    assert calls == []


# --- generate_package end-to-end --------------------------------------------


def test_generate_package_skips_generation_when_verdict_is_not_pursue(tmp_path):
    client = _FakeClient(responses=[(_EVAL_PAYLOAD_PASS, 900, 150)])
    result = llm_apply.generate_package(
        "JD text", company="Acme", title="Engineer", model="claude-haiku-4-5", client=client, output_root=tmp_path
    )

    assert result.evaluation.verdict == "pass"
    assert result.resume_path is None
    assert result.cover_letter_path is None
    assert result.generate_metrics is None
    assert len(client.calls) == 1  # only the evaluate call, no generation
    assert result.total_input_tokens == 900
    assert result.total_output_tokens == 150

    # The JD text + review are still written even on a non-pursue verdict.
    assert result.jd_path is not None and result.jd_path.exists()
    assert result.jd_path.name == "JobDescription.docx"
    assert result.review_path is not None and result.review_path.exists()
    assert result.review_path.name == "full-LLM-review.docx"
    # Both artifacts land together in one per-job folder. This is the only
    # tracked lead for "Acme" (no multi_lead passed), so the layout is flat:
    # straight in the company folder, no title-specific subfolder.
    assert result.jd_path.parent == result.review_path.parent == tmp_path / "Acme"

    review_doc = Document(result.review_path)
    review_text = "\n".join(cell.text for table in review_doc.tables for row in table.rows for cell in row.cells)
    assert "Angular required" in review_text


def test_generate_package_force_generates_despite_non_pursue_verdict(tmp_path):
    client = _FakeClient(
        responses=[
            (_EVAL_PAYLOAD_PASS, 900, 150),
            (json.dumps(_CLEAN_CONTENT), 2000, 3000),
        ]
    )
    result = llm_apply.generate_package(
        "JD text",
        company="Acme",
        title="Engineer",
        model="claude-haiku-4-5",
        client=client,
        output_root=tmp_path,
        force=True,
    )

    assert result.evaluation.verdict == "pass"
    assert result.resume_path is not None and result.resume_path.exists()
    assert result.cover_letter_path is not None and result.cover_letter_path.exists()
    assert result.generate_metrics is not None
    assert len(client.calls) == 2  # evaluate + generate, unlike the non-forced case


def test_generate_package_full_flow_saves_docx_and_aggregates_metrics(tmp_path):
    client = _FakeClient(
        responses=[
            (_EVAL_PAYLOAD_PURSUE, 1000, 200),
            (json.dumps(_CLEAN_CONTENT), 2000, 3000),
        ]
    )
    result = llm_apply.generate_package(
        "JD text", company="Bellese", title="Senior Engineer", model="claude-haiku-4-5", client=client, output_root=tmp_path
    )

    assert result.evaluation.verdict == "pursue"
    assert result.jd_path is not None and result.jd_path.exists()
    assert result.review_path is not None and result.review_path.exists()
    assert result.resume_path is not None and result.resume_path.exists()
    assert result.cover_letter_path is not None and result.cover_letter_path.exists()
    assert result.warnings == []
    # All four artifacts land in the same per-job folder. Flat layout: this
    # is the only tracked lead for "Bellese" (no multi_lead passed).
    job_folder = tmp_path / "Bellese"
    assert result.jd_path.parent == result.review_path.parent == job_folder
    assert result.resume_path.parent == result.cover_letter_path.parent == job_folder

    assert result.generate_metrics.input_tokens == 2000
    assert result.generate_metrics.output_tokens == 3000
    assert result.total_input_tokens == 1000 + 2000
    assert result.total_output_tokens == 200 + 3000
    expected_cost = (
        (1000 / 1_000_000 * 1.00 + 200 / 1_000_000 * 5.00)
        + (2000 / 1_000_000 * 1.00 + 3000 / 1_000_000 * 5.00)
    )
    assert result.total_cost_usd == pytest.approx(expected_cost)

    # HomePortfolio (listed first in _CLEAN_CONTENT) must be rendered last.
    doc = Document(result.resume_path)
    employer_lines = [p.text for p in doc.paragraphs if "HomePortfolio" in p.text or "Spexture" in p.text]
    assert employer_lines[-1].startswith("HomePortfolio")


def test_generate_package_multi_lead_nests_under_company_folder(tmp_path):
    client = _FakeClient(responses=[(_EVAL_PAYLOAD_PASS, 900, 150)])
    result = llm_apply.generate_package(
        "JD text",
        company="Acme",
        title="Engineer",
        model="claude-haiku-4-5",
        client=client,
        output_root=tmp_path,
        multi_lead=True,
        sibling_titles=("Manager",),
    )
    assert result.jd_path.parent == result.review_path.parent == tmp_path / "Acme" / "Acme_Engineer"


# --- _job_folder layout -------------------------------------------------------


def test_job_folder_flat_when_not_multi_lead(tmp_path):
    folder = llm_apply._job_folder(tmp_path, company="Acme", title="Engineer", multi_lead=False)
    assert folder == tmp_path / "Acme"


def test_job_folder_nested_when_multi_lead(tmp_path):
    folder = llm_apply._job_folder(tmp_path, company="Acme", title="Engineer", multi_lead=True)
    assert folder == tmp_path / "Acme" / "Acme_Engineer"


def test_job_folder_migrates_stale_flat_files_when_second_lead_appears(tmp_path):
    # Simulate a prior single-lead run that wrote flat files for "Manager".
    flat_dir = tmp_path / "Acme"
    flat_dir.mkdir()
    (flat_dir / "JobDescription.docx").write_text("old JD", encoding="utf-8")

    # A second lead ("Engineer") now shows up for the same company.
    folder = llm_apply._job_folder(
        tmp_path, company="Acme", title="Engineer", multi_lead=True, sibling_titles=("Manager",)
    )

    assert folder == flat_dir / "Acme_Engineer"
    migrated = flat_dir / "Acme_Manager" / "JobDescription.docx"
    assert migrated.exists()
    assert migrated.read_text(encoding="utf-8") == "old JD"
    assert not (flat_dir / "JobDescription.docx").exists()


def test_job_folder_does_not_guess_when_migration_is_ambiguous(tmp_path):
    flat_dir = tmp_path / "Acme"
    flat_dir.mkdir()
    (flat_dir / "JobDescription.docx").write_text("old JD", encoding="utf-8")

    # Two (or zero) apparent siblings -> ambiguous which lead the stray flat
    # files belong to; leave them alone rather than guessing wrong.
    folder = llm_apply._job_folder(
        tmp_path, company="Acme", title="Engineer", multi_lead=True, sibling_titles=("Manager", "Director")
    )

    assert folder == flat_dir / "Acme_Engineer"
    assert (flat_dir / "JobDescription.docx").exists()  # left in place, not moved


def test_generate_package_auto_repairs_house_rule_violation_before_saving(tmp_path):
    dirty_content = {
        "resume": _CLEAN_CONTENT["resume"],
        "cover_letter": {"salutation": "Dear Hiring Team,", "paragraphs": ["I am a US citizen and eligible for Public Trust clearance."]},
    }
    repaired_content = {
        "resume": _CLEAN_CONTENT["resume"],
        "cover_letter": {"salutation": "Dear Hiring Team,", "paragraphs": ["I bring deep relevant experience."]},
    }
    client = _FakeClient(
        responses=[
            (_EVAL_PAYLOAD_PURSUE, 1000, 200),
            (json.dumps(dirty_content), 1500, 1500),
            (json.dumps(repaired_content), 400, 300),
        ]
    )
    result = llm_apply.generate_package(
        "JD text", company="Bellese", title="Senior Engineer", model="claude-haiku-4-5", client=client, output_root=tmp_path
    )

    assert result.warnings == []  # cleared after the repair pass
    assert len(client.calls) == 3  # evaluate, generate, house-rule repair
    # Repair-call tokens must be folded into the generate step total.
    assert result.generate_metrics.input_tokens == 1500 + 400
    assert result.generate_metrics.output_tokens == 1500 + 300

    doc = Document(result.cover_letter_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "citizen" not in full_text.lower()
    assert "deep relevant experience" in full_text


def test_generate_package_leaves_warnings_if_repair_cannot_clear_them(tmp_path):
    dirty_content = {
        "resume": _CLEAN_CONTENT["resume"],
        "cover_letter": {"paragraphs": ["I am a US citizen."]},
    }
    still_dirty_content = {
        "resume": _CLEAN_CONTENT["resume"],
        "cover_letter": {"paragraphs": ["I am still a US citizen, for emphasis."]},
    }
    client = _FakeClient(
        responses=[
            (_EVAL_PAYLOAD_PURSUE, 1000, 200),
            (json.dumps(dirty_content), 1000, 1000),
            (json.dumps(still_dirty_content), 400, 300),
        ]
    )
    result = llm_apply.generate_package(
        "JD text", company="Bellese", title="Senior Engineer", model="claude-haiku-4-5", client=client, output_root=tmp_path
    )

    assert result.warnings  # repair pass didn't fully clear the violation
    assert result.resume_path is not None  # still saved — warnings are non-blocking


# --- cost/metrics helpers -----------------------------------------------------


def test_cost_usd_known_model():
    cost = llm_apply._cost_usd("claude-haiku-4-5", 1_000_000, 1_000_000)
    assert cost == pytest.approx(1.00 + 5.00)


def test_cost_usd_unknown_model_returns_none():
    assert llm_apply._cost_usd("some-future-model", 1000, 1000) is None


def test_sum_metrics_aggregates_multiple_calls():
    calls = [
        llm_apply.CallMetrics(step="generate", model="claude-haiku-4-5", input_tokens=100, output_tokens=50, elapsed_s=1.0, cost_usd=0.001),
        llm_apply.CallMetrics(step="generate_house_rule_repair", model="claude-haiku-4-5", input_tokens=40, output_tokens=20, elapsed_s=0.5, cost_usd=0.0005),
    ]
    total = llm_apply._sum_metrics("generate", "claude-haiku-4-5", calls)
    assert total.input_tokens == 140
    assert total.output_tokens == 70
    assert total.elapsed_s == pytest.approx(1.5)
    assert total.cost_usd == pytest.approx(0.0015)


def test_sum_metrics_handles_empty_call_list():
    total = llm_apply._sum_metrics("generate", "claude-haiku-4-5", [])
    assert total.input_tokens == 0
    assert total.cost_usd is None


# --- rendering helpers --------------------------------------------------------


def test_safe_filename_strips_unsafe_characters():
    assert llm_apply._safe_filename("Shawn Becker: Résumé (Draft)!.docx") == "Shawn_Becker_Rsum_Draft.docx"


def test_render_resume_places_homeportfolio_last_regardless_of_input_order(tmp_path):
    resume = {
        "experience": [
            {"employer": "HomePortfolio", "dates": "1997-2002", "bullets": ["a"]},
            {"employer": "Sierra Vista Group", "dates": "2002-2011", "bullets": ["b"]},
        ]
    }
    path = llm_apply.render_resume(resume, company="Acme", title="Engineer", out_dir=tmp_path)
    assert path.parent == tmp_path / "Acme"
    doc = Document(path)
    employer_paras = [p.text for p in doc.paragraphs if "HomePortfolio" in p.text or "Sierra Vista" in p.text]
    assert employer_paras[0].startswith("Sierra Vista")
    assert employer_paras[1].startswith("HomePortfolio")


def test_render_cover_letter_includes_phone_and_salutation(tmp_path):
    cover_letter = {"salutation": "Dear Hiring Team,", "paragraphs": ["Body paragraph."]}
    path = llm_apply.render_cover_letter(cover_letter, company="Acme", title="Engineer", out_dir=tmp_path)
    assert path.parent == tmp_path / "Acme"
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert llm_apply.CANDIDATE_PHONE in full_text
    assert "Dear Hiring Team," in full_text
    assert "Body paragraph." in full_text


def test_render_resume_and_cover_letter_share_the_same_job_folder(tmp_path):
    resume_path = llm_apply.render_resume({}, company="Acme", title="Engineer", out_dir=tmp_path)
    cover_letter_path = llm_apply.render_cover_letter({}, company="Acme", title="Engineer", out_dir=tmp_path)
    assert resume_path.parent == cover_letter_path.parent == tmp_path / "Acme"


def test_render_resume_and_cover_letter_nest_under_company_when_multi_lead(tmp_path):
    resume_path = llm_apply.render_resume(
        {}, company="Acme", title="Engineer", out_dir=tmp_path, multi_lead=True, sibling_titles=("Manager",)
    )
    cover_letter_path = llm_apply.render_cover_letter(
        {}, company="Acme", title="Engineer", out_dir=tmp_path, multi_lead=True, sibling_titles=("Manager",)
    )
    assert resume_path.parent == cover_letter_path.parent == tmp_path / "Acme" / "Acme_Engineer"


# --- JD review rendering -------------------------------------------------------


def _sample_evaluation(**overrides) -> llm_apply.EvaluationResult:
    defaults = dict(
        verdict="pursue",
        match_pct=92.0,
        job_summary="A greenfield agentic-AI team in regulated healthcare.",
        dealbreaker_checks=[
            {"check": "Banned stack", "status": "clean", "notes": "Python/TypeScript, React/Next.js."},
            {"check": "Comp floor", "status": "clean", "notes": "$170K-$195K, well above floor."},
        ],
        skills_alignment=[
            {"requirement": "Agentic systems", "evidence": "healthcare-agentic-snowflake-rag", "strength": "strong"},
            {"requirement": "React/Next.js", "evidence": "React yes; Next.js not in portfolio", "strength": "minor_gap"},
        ],
        flags=["Title reads Senior but requirements read mid-level."],
        rationale="Cleanest skill fit in recent pipeline.",
        framing_guidance=["Lead with the healthcare-agentic-RAG throughline."],
        structural_verdict="PASS on structure",
        next_step="",
        cover_letter_strategy="Lead with the healthcare-agentic-RAG throughline as the core narrative.",
        interview_prep=["Open with the healthcare-agentic-RAG project as the go-to example."],
    )
    defaults.update(overrides)
    return llm_apply.EvaluationResult(**defaults)


def test_render_jd_review_includes_all_sections():
    text = llm_apply.render_jd_review(_sample_evaluation(), company="Talkiatry", title="Senior AI Engineer")

    assert "Senior AI Engineer @ Talkiatry" in text
    assert "About the job" in text
    assert "greenfield agentic-AI team" in text
    assert "Dealbreaker sweep" in text
    assert "Banned stack" in text and "✅ Clean" in text
    assert "No hard dealbreakers." in text
    assert "Skills alignment" in text
    assert "Strong (real overlap):" in text
    assert "Agentic systems" in text and "healthcare-agentic-snowflake-rag" in text
    assert "Gaps:" in text and "React/Next.js" in text
    assert "Verdict: ~92% skills match / PASS on structure." in text
    assert "Flags" in text
    assert "seniority" in text.lower() or "mid-level" in text.lower()
    assert "Recommendation: PURSUE" in text
    assert "healthcare-agentic-RAG throughline" in text
    assert "### Cover letter strategy" in text
    assert "core narrative" in text
    assert "### Interview prep" in text
    assert "go-to example" in text


def test_render_jd_review_reports_fired_dealbreaker():
    evaluation = _sample_evaluation(
        verdict="pass",
        dealbreaker_checks=[{"check": "Banned stack", "status": "fail", "notes": "Angular required."}],
        flags=[],
        framing_guidance=[],
    )
    text = llm_apply.render_jd_review(evaluation, company="Acme", title="Engineer")
    assert "1 hard dealbreaker(s) fired." in text
    assert "Recommendation: PASS" in text
    # No Flags section when there are no flags.
    assert "### Flags" not in text


def test_render_jd_review_includes_next_step_escape_hatch_when_present():
    evaluation = _sample_evaluation(next_step="Ask the recruiter to confirm remote eligibility before applying.")
    text = llm_apply.render_jd_review(evaluation, company="Acme", title="Engineer")
    assert "**Next step:** Ask the recruiter to confirm remote eligibility before applying." in text


def test_render_jd_review_docx_writes_llm_review_with_dealbreaker_table_and_grouped_skills(tmp_path):
    path = llm_apply.render_jd_review_docx(_sample_evaluation(), company="Talkiatry", title="Senior AI Engineer", out_dir=tmp_path)
    assert path.exists()
    assert path.name == "full-LLM-review.docx"
    assert path.parent == tmp_path / "Talkiatry"

    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Senior AI Engineer @ Talkiatry" in full_text
    assert "greenfield agentic-AI team" in full_text
    assert "Verdict: ~92% skills match / PASS on structure." in full_text
    assert "Recommendation: PURSUE" in full_text
    assert "healthcare-agentic-RAG throughline" in full_text
    assert "Strong (real overlap):" in full_text
    assert "Agentic systems" in full_text and "healthcare-agentic-snowflake-rag" in full_text
    assert "Gaps: " in full_text and "React/Next.js" in full_text
    assert "Cover letter strategy" in full_text and "core narrative" in full_text
    assert "Interview prep" in full_text and "go-to example" in full_text

    # Skills alignment is grouped prose now, not a table — only the
    # dealbreaker sweep still renders as a real table.
    assert len(doc.tables) == 1
    dealbreaker_table_text = [cell.text for row in doc.tables[0].rows for cell in row.cells]
    assert "Banned stack" in dealbreaker_table_text
    assert "✅ Clean" in dealbreaker_table_text


def test_render_job_description_writes_docx(tmp_path):
    path = llm_apply.render_job_description(
        "We are hiring a Senior AI Engineer.\n\nResponsibilities:\n- Build things.",
        company="Talkiatry",
        title="Senior AI Engineer",
        out_dir=tmp_path,
    )
    assert path.exists()
    assert path.name == "JobDescription.docx"
    assert path.parent == tmp_path / "Talkiatry"

    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Senior AI Engineer @ Talkiatry" in full_text
    assert "Build things." in full_text


def test_job_description_and_review_share_the_same_job_folder(tmp_path):
    jd_path = llm_apply.render_job_description("JD text", company="Acme", title="Engineer", out_dir=tmp_path)
    review_path = llm_apply.render_jd_review_docx(_sample_evaluation(), company="Acme", title="Engineer", out_dir=tmp_path)
    assert jd_path.parent == review_path.parent == tmp_path / "Acme"


def test_job_description_and_review_nest_under_company_when_multi_lead(tmp_path):
    jd_path = llm_apply.render_job_description(
        "JD text", company="Acme", title="Engineer", out_dir=tmp_path, multi_lead=True, sibling_titles=("Manager",)
    )
    review_path = llm_apply.render_jd_review_docx(
        _sample_evaluation(), company="Acme", title="Engineer", out_dir=tmp_path,
        multi_lead=True, sibling_titles=("Manager",),
    )
    assert jd_path.parent == review_path.parent == tmp_path / "Acme" / "Acme_Engineer"


# --- candidate profile loading -------------------------------------------------


def test_load_candidate_profile_raises_when_missing(monkeypatch, tmp_path):
    monkeypatch.setattr(llm_apply, "_CANDIDATE_PROFILE_PATH", tmp_path / "does_not_exist.md")
    with pytest.raises(llm_apply.LLMApplyError):
        llm_apply._load_candidate_profile()
