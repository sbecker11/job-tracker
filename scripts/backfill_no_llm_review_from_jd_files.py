#!/usr/bin/env python3
"""Filesystem-driven complement to `backfill_jd_and_no_llm_review.py`.

That script works from the DB outward: for every *lead*, check its computed
folder for the two files. This script works from disk inward: find every
folder that already HAS a `JobDescription.docx` — regardless of whether it
has a matching DB row, correct naming, or was placed there by some older/
manual process — and make sure `no-LLM-review.docx` sits next to it.

Catches gaps the DB-driven pass can't see: legacy/orphan folders (e.g. a
JD saved under an old naming convention, or one with no corresponding
`normalized_key` at all after a rename/merge).

For each `JobDescription.docx` found with no sibling `no-LLM-review.docx`:
  1. Try to match its folder back to a DB lead (via `_safe_filename`
     comparison against every lead's computed folder) to get the canonical
     company/title and stored `jd_text`.
  2. If no DB match, fall back to reading the JD text directly out of the
     docx file itself (the heading paragraph `"{title} @ {company}"` plus
     the body paragraphs render_job_description wrote).
  3. Score it (free, rule-based, no LLM call) and write `no-LLM-review.docx`
     directly into the SAME folder the JD was found in — deliberately not
     recomputed via `_job_folder`, so this can never write to the "wrong"
     location for a legacy/oddly-named folder.

Usage:
    python scripts/backfill_no_llm_review_from_jd_files.py [--dry-run] [--db PATH] [--output-root PATH]
"""

from __future__ import annotations

import argparse
import sqlite3
import sys
from pathlib import Path

from docx import Document

_REPO_ROOT = Path(__file__).resolve().parents[1]
_SRC = _REPO_ROOT / "src"
if str(_SRC) not in sys.path:
    sys.path.insert(0, str(_SRC))

from job_tracker.pipeline.llm_apply import DEFAULT_OUTPUT_ROOT, _safe_filename  # noqa: E402
from job_tracker.pipeline.store import DEFAULT_DB_PATH, get_sibling_titles  # noqa: E402
from job_tracker.scoring.scorer import score_jd  # noqa: E402

# Mirrors render_no_llm_review_docx's own rendering, but targets an explicit
# already-known folder instead of recomputing one via _job_folder — see
# module docstring for why.
from job_tracker.pipeline.llm_apply import _RULE_STATUS_LABEL, _add_table, should_run_llm_review  # noqa: E402


def _lead_folder(out_dir: Path, *, company: str, title: str, multi_lead: bool) -> Path:
    company_dir = out_dir / _safe_filename(company)
    if not multi_lead:
        return company_dir
    return company_dir / _safe_filename(f"{company}_{title}")


def _extract_jd_from_docx(path: Path) -> tuple[str, str, str]:
    """Returns (company, title, jd_text) parsed back out of a
    render_job_description-style docx: first paragraph is the
    `"{title} @ {company}"` heading, the rest is the JD body."""
    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    company, title = "", ""
    body_start = 0
    if paras and " @ " in paras[0]:
        title, _, company = paras[0].rpartition(" @ ")
        body_start = 1
    jd_text = "\n\n".join(paras[body_start:])
    return company, title, jd_text


def _render_no_llm_review_to(out_path: Path, score, *, company: str, title: str) -> None:
    doc = Document()
    doc.add_heading(f"{title} @ {company}", level=1)
    p = doc.add_paragraph()
    p.add_run("Rule-based review — no LLM call made.").italic = True

    if score.dealbreaker_hits:
        doc.add_heading("Dealbreaker sweep (rule-based)", level=2)
        _add_table(
            doc,
            headers=("Check", "Status", "Hits"),
            rows=[
                (h.label, _RULE_STATUS_LABEL[h.load_bearing], str(h.hit_count))
                for h in score.dealbreaker_hits
            ],
        )
        n_fail = sum(1 for h in score.dealbreaker_hits if h.load_bearing)
        p = doc.add_paragraph()
        p.add_run("No hard dealbreakers." if n_fail == 0 else f"{n_fail} hard dealbreaker(s) fired.").bold = True

    doc.add_heading("Skills alignment (rule-based, keyword match)", level=2)
    if score.matched_skills:
        doc.add_paragraph("Matched: " + ", ".join(sorted(score.matched_skills)))
    else:
        doc.add_paragraph("No known skills matched against this JD text.")

    p = doc.add_paragraph()
    p.add_run(f"Match: ~{score.match_pct:.0f}% (rule-based, JD-relative).").bold = True
    p.add_run(f" Recognized JD tech vocabulary weight: {score.relevant_weight:.0f}.")

    if score.rationale:
        doc.add_heading("Rationale", level=2)
        for line in score.rationale:
            doc.add_paragraph(line, style="List Bullet")

    doc.add_heading(f"Recommendation: {score.verdict.upper()}", level=2)
    gate = should_run_llm_review(score)
    doc.add_paragraph(
        "This score clears the full-LLM-review threshold — a deeper LLM-backed review follows."
        if gate
        else "Below the full-LLM-review threshold — no LLM call made for this lead."
    )
    doc.save(str(out_path))


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--db", type=Path, default=DEFAULT_DB_PATH)
    ap.add_argument("--output-root", type=Path, default=DEFAULT_OUTPUT_ROOT)
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args(argv)

    conn = sqlite3.connect(args.db, timeout=30)
    conn.row_factory = sqlite3.Row
    leads = conn.execute("SELECT normalized_key, company, title, jd_text FROM job_leads").fetchall()

    # Build folder -> (company, title, jd_text) index for every lead, so a
    # found JobDescription.docx can be matched back to its canonical DB row.
    folder_index: dict[Path, sqlite3.Row] = {}
    for lead in leads:
        sibling_titles = tuple(get_sibling_titles(conn, lead["company"], exclude_title=lead["title"]))
        multi_lead = len(sibling_titles) > 0
        folder = _lead_folder(args.output_root, company=lead["company"], title=lead["title"], multi_lead=multi_lead)
        folder_index[folder.resolve()] = lead

    jd_files = sorted(args.output_root.rglob("JobDescription.docx"))
    print(f"found {len(jd_files)} JobDescription.docx file(s) under {args.output_root}")

    written = 0
    already_ok = 0
    matched_db = 0
    fell_back_to_docx = 0
    failures: list[tuple[Path, str]] = []

    for jd_path in jd_files:
        folder = jd_path.parent
        review_path = folder / "no-LLM-review.docx"
        if review_path.is_file():
            already_ok += 1
            continue

        lead = folder_index.get(folder.resolve())
        try:
            if lead is not None and lead["jd_text"]:
                company, title, jd_text = lead["company"], lead["title"], lead["jd_text"]
                matched_db += 1
            else:
                company, title, jd_text = _extract_jd_from_docx(jd_path)
                fell_back_to_docx += 1

            if len(jd_text or "") < 50:
                failures.append((folder, "extracted JD text too short to score"))
                continue

            score = score_jd(jd_text)
            print(f"{'[dry-run] would write' if args.dry_run else 'writing'} no-LLM-review.docx: {folder}")
            if not args.dry_run:
                _render_no_llm_review_to(review_path, score, company=company or folder.name, title=title or folder.name)
            written += 1
        except Exception as exc:  # noqa: BLE001 — best-effort repair pass, keep going on any one bad file
            failures.append((folder, str(exc)))

    print()
    print(f"already had no-LLM-review.docx: {already_ok}")
    print(f"{'would write' if args.dry_run else 'wrote'} no-LLM-review.docx: {written}  (matched to DB lead: {matched_db}, extracted from docx: {fell_back_to_docx})")
    print(f"failures: {len(failures)}")
    for folder, reason in failures:
        print(f"  {folder}: {reason}")

    conn.close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
