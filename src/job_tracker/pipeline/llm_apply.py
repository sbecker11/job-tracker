"""LLM-driven JD evaluation + résumé/cover-letter generation.

Runs the same "JD Match Framework" (dealbreaker sweep -> skills alignment ->
match % -> verdict) documented in ~/CLAUDE.md §10, and — only on a
"pursue" verdict — generates tailored résumé + cover letter content per
§5-§9 and §11, then renders it to real .docx files.

This is a plain two-stage function pipeline (evaluate, then generate), not a
service architecture: each stage is a single structured LLM call with a
fixed input/output shape, not an autonomous multi-step agent, so a
synchronous in-process call is the right level of complexity — see the
chat history for why this was deliberately NOT split into MCP/FastAPI
services.

Contact-info constants below are hardcoded from CLAUDE.md §1 rather than
left to the model to reproduce verbatim, since a hallucinated typo in an
email/phone number is a much worse failure mode than losing a little
creative freedom the model doesn't need for that text anyway.
"""

from __future__ import annotations

import json
import logging
import os
import re
import time
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from job_tracker.pipeline.letter_style import (
    add_muted_contact_line,
    apply_template_styles,
    looks_like_signature_line,
)
from job_tracker.scoring.scorer import DealbreakerHit, ScoreResult, score_jd, should_run_llm_review

logger = logging.getLogger(__name__)

DEFAULT_MODEL = os.environ.get("JOB_TRACKER_APPLY_MODEL", "claude-sonnet-5")

# USD per million tokens: (input, output). Source: platform.claude.com/docs/en/about-claude/pricing
# (checked 2026-07-03). Sonnet 5 is on introductory pricing through 2026-08-31, then $3/$15.
_MODEL_PRICING_USD_PER_MTOK: dict[str, tuple[float, float]] = {
    "claude-haiku-4-5": (1.00, 5.00),
    "claude-sonnet-4-5": (3.00, 15.00),
    "claude-sonnet-5": (2.00, 10.00),
    "claude-opus-4-8": (5.00, 25.00),
}

_CANDIDATE_PROFILE_PATH = Path(
    os.environ.get("JOB_TRACKER_CANDIDATE_PROFILE_PATH", str(Path.home() / "CLAUDE.md"))
)

DEFAULT_OUTPUT_ROOT = Path.home() / "Desktop" / "Resumes" / "2026"

CANDIDATE_NAME = "Shawn Becker"
CANDIDATE_EMAIL = "shawn.becker@spexture.com"
CANDIDATE_PHONE = "+1 857-891-0896"
CANDIDATE_LINKEDIN = "linkedin.com/in/shawnbecker"
CANDIDATE_GITHUB = "github.com/sbecker11"

# CLAUDE.md §4 "Banned terms" — checked mechanically post-generation as a
# safety net; the prompt already instructs the model to avoid these, but a
# hardcoded name-collision (e.g. "Cambria" vs "Cambia Health Solutions") is
# exactly the kind of thing worth a belt-and-suspenders check.
_BANNED_TERMS = ["Spexture LLC", "sub-100ms response", "Cline", "Member Nav", "Cambria"]

# CLAUDE.md §4 house rule #11 bans any work-authorization statement outright.
# This is a natural-language rule an LLM can quietly overstep (observed in
# testing: it added a citizenship/clearance statement to a cover letter for
# a role that didn't ask for one), so it gets a regex safety net on top of
# the prompt.
_WORK_AUTH_RE = re.compile(
    r"\bUS citizen(ship)?\b|\bgreen card\b|authorized to work|work authorization|"
    r"eligible for (a )?(public trust|security clearance|clearance)|"
    r"sponsorship (is )?(not )?(required|needed|available)",
    re.IGNORECASE,
)
# CLAUDE.md §4 house rule #12 (added 2026-07-05 after a real leak was found in
# the corpus review: a cover letter compared a W2 rate against an equivalent
# C2C rate). No dollar figure, hourly rate, or salary range belongs in either
# deliverable — compensation is a live conversation, never written down here.
_COMP_FIGURE_RE = re.compile(
    r"\$\s?\d[\d,.]*\s*(?:/\s*hr|per\s*hour|/\s*hour|k\b)|"
    r"\b\d[\d,.]*\s*(?:k|K)\s*(?:/\s*yr|per\s*year|base|salary)|"
    r"\bhourly rate\b|\bsalary range\b|\bcompensation range\b",
    re.IGNORECASE,
)

_UNSAFE_FILENAME_CHARS = re.compile(r"[^A-Za-z0-9_.-]")


class LLMApplyError(RuntimeError):
    """Raised when the candidate profile is missing or the LLM call fails unrecoverably."""


@dataclass
class CallMetrics:
    """Usage/cost accounting for one Anthropic API call."""

    step: str  # "evaluate" | "generate" | "house_rule_repair"
    model: str
    input_tokens: int = 0
    output_tokens: int = 0
    elapsed_s: float = 0.0
    cost_usd: float | None = None


def _cost_usd(model: str, input_tokens: int, output_tokens: int) -> float | None:
    pricing = _MODEL_PRICING_USD_PER_MTOK.get(model)
    if pricing is None:
        return None
    in_rate, out_rate = pricing
    return (input_tokens / 1_000_000) * in_rate + (output_tokens / 1_000_000) * out_rate


def _sum_metrics(step: str, model: str, calls: list[CallMetrics]) -> CallMetrics:
    """Roll up every call made in service of one logical step (including any
    JSON-repair or house-rule-repair retries) into a single step-level total,
    so "cost of the generate step" means everything it took to get a usable
    result, not just the first (possibly-failed) call."""
    input_tokens = sum(c.input_tokens for c in calls)
    output_tokens = sum(c.output_tokens for c in calls)
    elapsed_s = sum(c.elapsed_s for c in calls)
    costs = [c.cost_usd for c in calls if c.cost_usd is not None]
    cost_usd = sum(costs) if costs else None
    return CallMetrics(
        step=step, model=model, input_tokens=input_tokens, output_tokens=output_tokens,
        elapsed_s=elapsed_s, cost_usd=cost_usd,
    )


@dataclass
class EvaluationResult:
    verdict: str  # "pursue" | "review" | "pass"
    match_pct: float
    job_summary: str = ""
    # Each item: {"check": str, "status": "clean" | "warning" | "fail", "notes": str}.
    # One row per CLAUDE.md §3 dealbreaker plus comp floor / remote-location fit /
    # durability lens (§2) — the model derives the exact set from whatever the
    # loaded candidate profile's dealbreaker table actually contains, so this
    # isn't hardcoded to today's five checks.
    dealbreaker_checks: list[dict] = field(default_factory=list)
    # Each item: {"requirement": str, "evidence": str, "strength": "very_strong" |
    # "strong" | "moderate" | "minor_gap" | "gap"}.
    skills_alignment: list[dict] = field(default_factory=list)
    # Notable non-dealbreaker concerns worth flagging (e.g. a title/seniority
    # mismatch, an ambiguous hybrid policy) — things that don't fail the
    # dealbreaker sweep but should still shape how the candidate approaches it.
    flags: list[str] = field(default_factory=list)
    rationale: str = ""
    # Only on pursue/review — concrete guidance for how to position the
    # application (what to lead with, what to downplay, an angle for the
    # cover letter) given whatever `flags` surfaced. Empty on a clean pass.
    framing_guidance: list[str] = field(default_factory=list)
    # CLAUDE.md §10 step 4 — a short skills-only structural read, independent
    # of whether dealbreakers actually fired, e.g. "PASS on structure" or
    # "FAIL on structure". Lets the verdict line separate "does this look
    # good on paper" from the final pursue/review/pass call.
    structural_verdict: str = ""
    # CLAUDE.md §10 step 5 — only when a dealbreaker is soft/confirmable
    # rather than absolute (e.g. remote eligibility unstated, not ruled
    # out): the one concrete, low-effort action that would resolve it
    # before writing the role off. Empty when nothing is actionable.
    next_step: str = ""
    # CLAUDE.md §10 step 6 — only on pursue/review: a short closing
    # paragraph (2-4 sentences) synthesizing framing_guidance into one
    # coherent cover-letter narrative angle, not a list of separate tactics.
    # Empty on a clean pass.
    cover_letter_strategy: str = ""
    # CLAUDE.md §10 step 7 — only on pursue/review: concrete talking points
    # for presenting this background out loud in an interview — what to
    # lead with verbally, how to preempt any flags, which specific
    # engagement/project to have ready as the go-to example. Empty on pass.
    interview_prep: list[str] = field(default_factory=list)
    metrics: CallMetrics | None = None


@dataclass
class PackageResult:
    evaluation: EvaluationResult
    jd_path: Path | None = None
    review_path: Path | None = None
    resume_path: Path | None = None
    cover_letter_path: Path | None = None
    warnings: list[str] = field(default_factory=list)
    generate_metrics: CallMetrics | None = None

    @property
    def total_input_tokens(self) -> int:
        return (self.evaluation.metrics.input_tokens if self.evaluation.metrics else 0) + (
            self.generate_metrics.input_tokens if self.generate_metrics else 0
        )

    @property
    def total_output_tokens(self) -> int:
        return (self.evaluation.metrics.output_tokens if self.evaluation.metrics else 0) + (
            self.generate_metrics.output_tokens if self.generate_metrics else 0
        )

    @property
    def total_elapsed_s(self) -> float:
        return (self.evaluation.metrics.elapsed_s if self.evaluation.metrics else 0.0) + (
            self.generate_metrics.elapsed_s if self.generate_metrics else 0.0
        )

    @property
    def total_cost_usd(self) -> float | None:
        parts = [
            m.cost_usd
            for m in (self.evaluation.metrics, self.generate_metrics)
            if m is not None and m.cost_usd is not None
        ]
        return sum(parts) if parts else None


def _load_candidate_profile() -> str:
    if not _CANDIDATE_PROFILE_PATH.exists():
        raise LLMApplyError(
            f"Candidate profile not found at {_CANDIDATE_PROFILE_PATH}. Set "
            "JOB_TRACKER_CANDIDATE_PROFILE_PATH if it lives somewhere else."
        )
    return _CANDIDATE_PROFILE_PATH.read_text(encoding="utf-8")


# The SDK's default is 600s per attempt (with retries on top) — fine for a
# one-off call, but a batch run (e.g. triage_recruiter_inbox.py over a large
# backlog) can silently stall for an hour-plus if even one call/retry hits a
# slow patch, since nothing in this codebase's call sites has its own
# timeout. Real evaluate/generate calls finish in well under a minute; a
# call still running after 2 is almost certainly stuck, not just slow, and
# job-tracker's built-in JSON-repair retry (see `_call_and_parse_json`)
# already covers the "got a bad response, try once more" case on top of this.
_ANTHROPIC_TIMEOUT_S = 120.0


def _client():
    import anthropic

    api_key = os.environ.get("ANTHROPIC_API_KEY")  # pragma: allowlist secret
    if not api_key:
        raise LLMApplyError("ANTHROPIC_API_KEY is not set (see .env.example).")
    return anthropic.Anthropic(api_key=api_key, timeout=_ANTHROPIC_TIMEOUT_S)  # pragma: allowlist secret


def _parse_json_object(raw: str) -> dict:
    text = raw.strip()
    if text.startswith("```"):
        text = text.strip("`").strip()
        if text.lower().startswith("json"):
            text = text[4:].strip()
    data = json.loads(text)
    if not isinstance(data, dict):
        raise LLMApplyError(f"expected a JSON object, got {type(data).__name__}")
    return data


def _estimate_tokens(text: str) -> int:
    """Rough English token estimate (~4 chars/token). Good enough for pre-call cost."""
    return max(1, (len(text) + 3) // 4)


def _fmt_usd(cost: float | None) -> str:
    return f"${cost:.4f}" if cost is not None else "n/a"


def _predicted_call_cost(
    model: str, system: str, user: str, *, max_tokens: int
) -> tuple[float | None, int, int]:
    """Pre-call cost estimate: chars/4 for input, ~25% of max_tokens for expected output."""
    est_in = _estimate_tokens(system) + _estimate_tokens(user)
    est_out = max(64, min(max_tokens, max_tokens // 4))
    return _cost_usd(model, est_in, est_out), est_in, est_out


def _print_llm_call_pred(step: str, cost: float | None, est_in: int, est_out: int) -> None:
    print(
        f"    [llm {step}] pred ~{_fmt_usd(cost)} (est. {est_in} in / ~{est_out} out)",
        flush=True,
    )


def _print_llm_call_actual(metrics: CallMetrics) -> None:
    """Emit approximate USD cost immediately after every Anthropic API call."""
    print(
        f"    [llm {metrics.step}] actual ~{_fmt_usd(metrics.cost_usd)} "
        f"({metrics.input_tokens} in / {metrics.output_tokens} out, {metrics.elapsed_s:.1f}s)",
        flush=True,
    )


def _call(system: str, user: str, *, model: str, client=None, max_tokens: int = 4096, step: str = "call") -> tuple[str, CallMetrics]:
    client = client or _client()
    pred_cost, est_in, est_out = _predicted_call_cost(model, system, user, max_tokens=max_tokens)
    _print_llm_call_pred(step, pred_cost, est_in, est_out)
    start = time.monotonic()
    response = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": user}],
    )
    elapsed_s = time.monotonic() - start
    text = "".join(block.text for block in response.content if getattr(block, "type", "") == "text")
    usage = getattr(response, "usage", None)
    input_tokens = getattr(usage, "input_tokens", 0) or 0
    output_tokens = getattr(usage, "output_tokens", 0) or 0
    metrics = CallMetrics(
        step=step, model=model, input_tokens=input_tokens, output_tokens=output_tokens,
        elapsed_s=elapsed_s, cost_usd=_cost_usd(model, input_tokens, output_tokens),
    )
    _print_llm_call_actual(metrics)
    return text, metrics


_REPAIR_SYSTEM_PROMPT = (
    "The text below was supposed to be a single valid JSON object but failed to parse "
    f"({{error}}). Output ONLY the corrected, valid JSON object — same content and keys, "
    "just fix the syntax (e.g. escape stray quotes inside string values). No markdown fences, no prose."
)


def _call_and_parse_json(
    system: str, user: str, *, model: str, client=None, max_tokens: int = 4096, step: str = "call"
) -> tuple[dict, list[CallMetrics]]:
    """Call the model and parse its JSON response, with one repair retry if the
    first response isn't valid JSON (observed failure mode: an unescaped quote
    inside a generated string value, e.g. around a quoted program name).
    Returns the parsed dict plus every CallMetrics incurred (1 normally, 2 if
    a repair retry was needed)."""
    client = client or _client()
    raw, metrics = _call(system, user, model=model, client=client, max_tokens=max_tokens, step=step)
    try:
        return _parse_json_object(raw), [metrics]
    except (json.JSONDecodeError, LLMApplyError) as exc:
        logger.warning("JSON parse failed, retrying with a repair call: %s", exc)
        repaired, repair_metrics = _call(
            _REPAIR_SYSTEM_PROMPT.format(error=exc),
            raw,
            model=model,
            client=client,
            max_tokens=max_tokens,
            step=f"{step}_json_repair",
        )
        return _parse_json_object(repaired), [metrics, repair_metrics]


_EVAL_SYSTEM_PROMPT = """You evaluate a job description against a specific candidate's real background, \
using the candidate profile document below, which defines the "JD Match Framework" the candidate uses \
(see its "JD Match Framework" section). Follow that framework's steps exactly: dealbreaker sweep, then \
skills alignment, then an honest match percentage, then a verdict. Produce a review detailed enough that \
the candidate can decide whether to pursue this role WITHOUT re-reading the original JD.

Rules:
- Never invent experience, skills, or history not present in the candidate profile.
- Use "pursue" when the dealbreaker sweep clears and the match is genuinely strong.
- Use "pass" when a real (load-bearing) dealbreaker fires, or the match is weak.
- Use "review" only for genuinely borderline cases worth a human's own judgment.
- Work authorization / citizenship (CRITICAL — evaluation vs deliverables): the candidate \
profile states (for targeting/evaluation only) that Shawn is a US citizen and needs no visa \
sponsorship. Treat JD requirements like 'US citizens only', 'must be a US citizen', 'authorized \
to work in the US without sponsorship', or equivalent as ✅ clean / a clear FIT in \
"dealbreaker_checks" — never "warning", never "fail", and never invent a next_step to \
"confirm citizenship". Do NOT choose verdict "review" solely because the JD requires \
citizenship or no-sponsorship work auth, and do NOT treat the deliverables ban on stating \
citizenship (house rule — résumé/cover letter must omit the topic) as an inability to \
confirm fit during evaluation. (Active security-clearance possession is a separate question \
from citizenship/sponsorship — only citizenship / no-sponsorship is auto-cleared here.)
- A skill or technology mentioned only as one alternative among several acceptable options \
(e.g. "React (or Angular)", "Golang, Java, Python, Ruby, C#, or similar") is NOT load-bearing on its own \
unless the surrounding text makes clear it's the team's actual primary/required choice.
- "job_summary": 2-3 plain sentences on what the role/team actually does, based on the JD text — not a \
sales pitch, just enough for the candidate to recall this specific posting later without re-reading it.
- "dealbreaker_checks": one row per check the candidate profile's dealbreaker table/targeting section \
defines (e.g. banned stack items, employment type, compensation floor, location/remote fit) PLUS the \
profile's "durability lens" if it has one. Each row's "status" is "clean" (no concern), "warning" (a \
soft/non-fatal concern worth noting), or "fail" (an actual dealbreaker). "notes" must be a SPECIFIC, \
narrative explanation that cites the JD's actual language for that check (quote or closely paraphrase the \
relevant phrase) — never a generic restatement of the rule with no citation. When citing the JD's exact \
wording, wrap it in single quotes ('...'), never double quotes.
- "skills_alignment": one row per significant JD requirement, each mapped to the candidate's actual \
matching evidence (a named real engagement or portfolio project from the profile, not a vague yes/no). \
"strength" is one of "very_strong", "strong", "moderate", "minor_gap", "gap" — "gap" only if the profile \
has genuinely nothing relevant to point to for that requirement. This is rendered later as a "Strong (real \
overlap):" list (requirement -> evidence, for very_strong/strong/moderate rows) followed by a single \
"Gaps:" line (for minor_gap/gap rows), so write each "evidence" string as a complete, standalone clause \
that reads well in that format (e.g. 'linkage-engine (Java 21, Spring Boot, Aurora PostgreSQL, ECS \
Fargate)'), not a sentence fragment that depends on surrounding table context.
- "flags": notable concerns that are NOT dealbreakers but still change how the candidate should approach \
this application — e.g. a title/seniority mismatch (JD's actual day-to-day responsibilities read as more \
junior or more senior than the title suggests), an unusually vague scope, a signal the team is much \
earlier/later stage than the durability lens prefers. Empty list if nothing stands out.
- "structural_verdict": a short skills-only structural read, e.g. "PASS on structure" or "FAIL on \
structure" — does this look like a good fit purely on skills/experience, independent of whether any \
dealbreaker actually fired. This can (and often will) differ from "verdict" once dealbreakers are factored \
in — e.g. a strong skills match that still gets a "pass" verdict because of a hard location dealbreaker.
- "next_step": only when a dealbreaker is soft/confirmable rather than absolute (e.g. remote eligibility \
is unstated in the JD rather than explicitly ruled out, or a compensation range is unlisted) — name the \
ONE concrete, low-effort action that would resolve that open question before writing the role off (e.g. \
"ask the recruiter whether this specific req is remote-eligible before investing in a full application"). \
Empty string when there is nothing actionable to check (e.g. the dealbreaker is stated outright, or there \
are no dealbreaker concerns at all).
- "framing_guidance": only populate on "pursue" or "review" — concrete, specific advice for HOW to \
position the application given whatever "flags" surfaced (what to lead with, what to downplay or reframe, \
a concrete cover-letter angle). Not generic advice ("highlight your skills") — it must reference the \
specific flags/gaps found for this role. Empty list on "pass".
- "cover_letter_strategy": only on "pursue" or "review" — a short closing paragraph (2-4 sentences) that \
synthesizes "framing_guidance" into ONE coherent narrative angle for the cover letter (not a restatement of \
the bullet list). Empty string on "pass".
- "interview_prep": only on "pursue" or "review" — a list of concrete talking points for presenting this \
background out loud in an interview: what to lead with verbally, how to preempt/address any "flags" above, \
and which specific named engagement/project to have ready as the go-to example for this role's core \
requirement. Not generic advice. Empty list on "pass".

JSON output requirements — read carefully, your response is parsed by a strict JSON parser with no tolerance \
for the mistakes below:
- Every string value must contain ONLY single quotes for any quoting — NEVER a double-quote character (") \
anywhere inside a string value. Double quotes are reserved for JSON syntax itself; one stray unescaped " \
inside a value breaks the entire response.
- Every string value must be a single line: never put a literal line break inside a string. If a note needs \
multiple sentences, separate them with a space, not a newline.
- Keep every string field concise (1-3 sentences, or one short clause for "evidence"/"requirement") — this \
is a structured record for quick scanning, not an essay.

Respond with ONLY a raw JSON object (no markdown fences, no prose outside the JSON), with exactly these keys:
{
  "job_summary": string,
  "dealbreaker_checks": [ {"check": string, "status": "clean" | "warning" | "fail", "notes": string}, ... ],
  "skills_alignment": [
    {"requirement": string, "evidence": string, "strength": "very_strong" | "strong" | "moderate" | "minor_gap" | "gap"}, ...
  ],
  "match_pct": number (0-100),
  "flags": [string, ...],
  "structural_verdict": string,
  "next_step": string,
  "verdict": "pursue" | "review" | "pass",
  "rationale": string,
  "framing_guidance": [string, ...],
  "cover_letter_strategy": string,
  "interview_prep": [string, ...]
}

--- CANDIDATE PROFILE ---
{profile}
--- END CANDIDATE PROFILE ---
"""

_GENERATE_SYSTEM_PROMPT = """You write a tailored résumé and cover letter for a candidate applying to a \
specific job, using ONLY the real background documented in the candidate profile below. Follow every \
house rule, structure rule, and content rule in that document exactly (canonical résumé structure, career \
timeline/dates, education formatting, banned terms, IC-only framing, and the "Generation Workflow & Output \
Conventions" content rules) — especially:
- Do NOT invent experience, employers, projects, or skills not present in the profile.
- Do NOT include any compensation figure/range, availability statement, or work-authorization statement of \
any kind — this includes citizenship, residency, green-card, sponsorship, or security-clearance-eligibility \
statements, EVEN IF the job description requires them or mentions clearance/citizenship requirements. Simply \
omit the topic entirely; do not reassure the reader about it.
- Do NOT include years on any degree, patent, or certification.
- Tailor bullets/skills selection to this specific JD by choosing from the candidate's real portfolio \
projects and technical anchors — pick what's actually relevant, don't list everything.
- If HomePortfolio is included in "experience", it must be the LAST entry.
- The cover letter body should tie 2-3 concrete real engagements/projects to this JD's stated needs.
- Do not write a salutation naming a specific person unless one is explicitly given in the job description \
below; otherwise use "Dear Hiring Team,".

Respond with ONLY a raw JSON object (no markdown fences, no prose outside the JSON), with exactly this shape:
{
  "resume": {
    "positioning_line": string,
    "summary": string,
    "skills": [string, ...],
    "experience": [
      {
        "employer": string,
        "dates": string,
        "role_note": string | null,
        "bullets": [string, ...],
        "subsections": [ {"heading": string, "bullets": [string, ...]}, ... ]
      }
    ],
    "education": [string, ...]
  },
  "cover_letter": {
    "salutation": string,
    "paragraphs": [string, ...]
  }
}

--- CANDIDATE PROFILE ---
{profile}
--- END CANDIDATE PROFILE ---
"""


def evaluate_lead(
    jd_text: str,
    *,
    company: str,
    title: str,
    model: str = DEFAULT_MODEL,
    client=None,
) -> EvaluationResult:
    profile = _load_candidate_profile()
    system = _EVAL_SYSTEM_PROMPT.replace("{profile}", profile)
    user = f"Company: {company}\nTitle: {title}\n\nJob description:\n{jd_text}"
    # max_tokens=16000 (raised from 8192 on 2026-07-12) — a batch run showed
    # ~25% of "evaluate" calls still truncating mid-JSON at 8192, always in
    # the schema's last fields (cover_letter_strategy/interview_prep) rather
    # than from a malformed/unescaped string, per direct inspection of the
    # raw (pre-parse) response text. The narrative-heavy 2026-07-11 schema
    # (job_summary, cited dealbreaker notes, per-row skills evidence,
    # framing_guidance/cover_letter_strategy/interview_prep) can genuinely
    # run past 8192 output tokens for a JD with many distinct requirements;
    # raising the ceiling doesn't raise cost since billing is by tokens
    # actually generated, not the ceiling.
    data, calls = _call_and_parse_json(system, user, model=model, client=client, max_tokens=16000, step="evaluate")

    def _dict_list(key: str, expected_keys: tuple[str, ...]) -> list[dict]:
        items = data.get(key) or []
        out = []
        for item in items:
            if isinstance(item, dict):
                out.append({k: str(item.get(k, "")) for k in expected_keys})
            elif item:
                # Defensive: tolerate an older/degraded flat-string response
                # (e.g. from a JSON-repair retry that lost structure) rather
                # than dropping the note entirely.
                out.append({expected_keys[0]: str(item)})
        return out

    return EvaluationResult(
        verdict=str(data.get("verdict", "review")).strip().lower(),
        match_pct=float(data.get("match_pct", 0) or 0),
        job_summary=str(data.get("job_summary", "")),
        dealbreaker_checks=_dict_list("dealbreaker_checks", ("check", "status", "notes")),
        skills_alignment=_dict_list("skills_alignment", ("requirement", "evidence", "strength")),
        flags=[str(f) for f in (data.get("flags") or [])],
        rationale=str(data.get("rationale", "")),
        framing_guidance=[str(g) for g in (data.get("framing_guidance") or [])],
        structural_verdict=str(data.get("structural_verdict", "")),
        next_step=str(data.get("next_step", "")),
        cover_letter_strategy=str(data.get("cover_letter_strategy", "")),
        interview_prep=[str(g) for g in (data.get("interview_prep") or [])],
        metrics=_sum_metrics("evaluate", model, calls),
    )


_STATUS_LABEL = {"clean": "✅ Clean", "warning": "⚠️ Warning", "fail": "🔴 Fail"}
_STRENGTH_LABEL = {
    "very_strong": "Very strong",
    "strong": "Strong",
    "moderate": "Moderate",
    "minor_gap": "Minor gap",
    "gap": "Gap",
}
_GAP_STRENGTHS = {"minor_gap", "gap"}


def _group_skills_alignment(skills_alignment: list[dict]) -> tuple[list[tuple[str, str]], str]:
    """Split skills_alignment rows into (a) "Strong (real overlap)" pairs of
    (requirement, evidence) for very_strong/strong/moderate rows, and (b) a
    single comma-joined "Gaps" summary string for minor_gap/gap rows — the
    grouped-prose format CLAUDE.md §10 step 2 specifies, mirroring how a
    live Claude-on-Chrome JD review presents this (table cells don't carry
    that narrative style as well as prose does)."""
    strong: list[tuple[str, str]] = []
    gap_parts: list[str] = []
    for row in skills_alignment:
        requirement = row.get("requirement", "")
        evidence = row.get("evidence", "")
        if row.get("strength", "") in _GAP_STRENGTHS:
            gap_parts.append(f"{requirement} ({evidence})" if evidence else requirement)
        else:
            strong.append((requirement, evidence))
    return strong, ", ".join(gap_parts)


def render_jd_review(evaluation: EvaluationResult, *, company: str, title: str) -> str:
    """Render an `EvaluationResult` as a human-readable markdown JD review —
    the "About the job / dealbreaker sweep / skills alignment / recommendation"
    report format the candidate reviews leads in (see chat history for the
    reference example this mirrors). Deterministic and free — pure formatting
    over already-computed evaluation data, no LLM call."""
    lines = [f"## {title} @ {company}", ""]

    if evaluation.job_summary:
        lines += ["### About the job", "", evaluation.job_summary, ""]

    if evaluation.dealbreaker_checks:
        lines += ["### Dealbreaker sweep", "", "| Check | Status | Notes |", "|---|---|---|"]
        for row in evaluation.dealbreaker_checks:
            status = _STATUS_LABEL.get(row.get("status", ""), row.get("status", ""))
            lines.append(f"| {row.get('check', '')} | {status} | {row.get('notes', '')} |")
        lines.append("")
        n_fail = sum(1 for r in evaluation.dealbreaker_checks if r.get("status") == "fail")
        lines.append("**No hard dealbreakers.**" if n_fail == 0 else f"**{n_fail} hard dealbreaker(s) fired.**")
        lines.append("")

    if evaluation.skills_alignment:
        strong, gaps = _group_skills_alignment(evaluation.skills_alignment)
        lines += ["### Skills alignment", ""]
        if strong:
            lines.append("**Strong (real overlap):**")
            lines += [f"- {req} → {ev}" if ev else f"- {req}" for req, ev in strong]
            lines.append("")
        if gaps:
            lines.append(f"**Gaps:** {gaps}")
            lines.append("")

    verdict_line = f"**Verdict: ~{evaluation.match_pct:.0f}% skills match"
    if evaluation.structural_verdict:
        verdict_line += f" / {evaluation.structural_verdict}"
    verdict_line += ".**"
    if evaluation.rationale:
        verdict_line += f" {evaluation.rationale}"
    lines.append(verdict_line.strip())
    lines.append("")

    if evaluation.next_step:
        lines.append(f"**Next step:** {evaluation.next_step}")
        lines.append("")

    if evaluation.flags:
        lines += ["### Flags", ""]
        lines += [f"- {flag}" for flag in evaluation.flags]
        lines.append("")

    lines.append(f"### Recommendation: {evaluation.verdict.upper()}")
    if evaluation.framing_guidance:
        lines.append("")
        lines += [f"- {g}" for g in evaluation.framing_guidance]
    lines.append("")

    if evaluation.cover_letter_strategy:
        lines += ["### Cover letter strategy", "", evaluation.cover_letter_strategy, ""]

    if evaluation.interview_prep:
        lines += ["### Interview prep", ""]
        lines += [f"- {tip}" for tip in evaluation.interview_prep]
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def _add_table(doc: Document, *, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value


_RULE_STATUS_LABEL = {True: "🔴 Fired (load-bearing)", False: "⚠️ Mentioned (not load-bearing)"}


def render_no_llm_review_docx(
    score: ScoreResult,
    *,
    company: str,
    title: str,
    out_dir: Path = DEFAULT_OUTPUT_ROOT,
    multi_lead: bool = False,
    sibling_titles: tuple[str, ...] = (),
) -> Path:
    """Render a `ScoreResult` (scoring.scorer.score_jd — free, deterministic,
    no LLM call) as `no-LLM-review.docx`: the first tier of the two-tier
    review pipeline (2026-07-11). Every lead gets one of these; only leads
    that clear `should_run_llm_review(score)` go on to a second,
    LLM-backed `full-LLM-review.docx` (see `render_jd_review_docx`)."""
    doc = Document()
    doc.add_heading(f"{title} @ {company}", level=1)
    p = doc.add_paragraph()
    p.add_run("Rule-based review — no LLM call made.").italic = True

    if score.dealbreaker_hits:
        doc.add_heading("Dealbreaker sweep (rule-based)", level=2)
        _add_table(
            doc,
            headers=("Check", "Status", "Hits"),
            rows=[
                (h.label, _RULE_STATUS_LABEL[h.load_bearing], str(h.hit_count))
                for h in score.dealbreaker_hits
            ],
        )
        n_fail = sum(1 for h in score.dealbreaker_hits if h.load_bearing)
        p = doc.add_paragraph()
        p.add_run("No hard dealbreakers." if n_fail == 0 else f"{n_fail} hard dealbreaker(s) fired.").bold = True

    doc.add_heading("Skills alignment (rule-based, keyword match)", level=2)
    if score.matched_skills:
        doc.add_paragraph("Matched: " + ", ".join(sorted(score.matched_skills)))
    else:
        doc.add_paragraph("No known skills matched against this JD text.")

    p = doc.add_paragraph()
    p.add_run(f"Match: ~{score.match_pct:.0f}% (rule-based, JD-relative).").bold = True
    p.add_run(f" Recognized JD tech vocabulary weight: {score.relevant_weight:.0f}.")

    if score.rationale:
        doc.add_heading("Rationale", level=2)
        for line in score.rationale:
            doc.add_paragraph(line, style="List Bullet")

    doc.add_heading(f"Recommendation: {score.verdict.upper()}", level=2)
    gate = should_run_llm_review(score)
    doc.add_paragraph(
        "This score clears the full-LLM-review threshold — a deeper LLM-backed review follows."
        if gate
        else "Below the full-LLM-review threshold — no LLM call made for this lead."
    )

    out_path = (
        _job_folder(out_dir, company=company, title=title, multi_lead=multi_lead, sibling_titles=sibling_titles)
        / "no-LLM-review.docx"
    )
    doc.save(str(out_path))
    return out_path


def render_jd_review_docx(
    evaluation: EvaluationResult,
    *,
    company: str,
    title: str,
    out_dir: Path = DEFAULT_OUTPUT_ROOT,
    multi_lead: bool = False,
    sibling_titles: tuple[str, ...] = (),
) -> Path:
    """Render an `EvaluationResult` as `full-LLM-review.docx` in that job's
    folder — the docx counterpart of `render_jd_review()`'s markdown, with
    the dealbreaker sweep and skills alignment as real tables instead of
    markdown pipe-tables (this is a saved artifact, not a terminal print).

    2026-07-11 two-tier restructure: this is now specifically the *second*
    tier (LLM-backed, nuanced) of the review pipeline — see
    `render_no_llm_review_docx` for the free/deterministic first tier that
    gates whether this one ever runs (`scoring.scorer.should_run_llm_review`)."""
    doc = Document()
    doc.add_heading(f"{title} @ {company}", level=1)

    if evaluation.job_summary:
        doc.add_heading("About the job", level=2)
        doc.add_paragraph(evaluation.job_summary)

    if evaluation.dealbreaker_checks:
        doc.add_heading("Dealbreaker sweep", level=2)
        _add_table(
            doc,
            headers=("Check", "Status", "Notes"),
            rows=[
                (
                    row.get("check", ""),
                    _STATUS_LABEL.get(row.get("status", ""), row.get("status", "")),
                    row.get("notes", ""),
                )
                for row in evaluation.dealbreaker_checks
            ],
        )
        n_fail = sum(1 for r in evaluation.dealbreaker_checks if r.get("status") == "fail")
        p = doc.add_paragraph()
        p.add_run(
            "No hard dealbreakers." if n_fail == 0 else f"{n_fail} hard dealbreaker(s) fired."
        ).bold = True

    if evaluation.skills_alignment:
        doc.add_heading("Skills alignment", level=2)
        strong, gaps = _group_skills_alignment(evaluation.skills_alignment)
        if strong:
            p = doc.add_paragraph()
            p.add_run("Strong (real overlap):").bold = True
            for req, ev in strong:
                doc.add_paragraph(f"{req} → {ev}" if ev else req, style="List Bullet")
        if gaps:
            p = doc.add_paragraph()
            p.add_run("Gaps: ").bold = True
            p.add_run(gaps)

    p = doc.add_paragraph()
    verdict_run_text = f"Verdict: ~{evaluation.match_pct:.0f}% skills match"
    if evaluation.structural_verdict:
        verdict_run_text += f" / {evaluation.structural_verdict}"
    verdict_run_text += "."
    p.add_run(verdict_run_text).bold = True
    if evaluation.rationale:
        p.add_run(f" {evaluation.rationale}")

    if evaluation.next_step:
        p = doc.add_paragraph()
        p.add_run("Next step: ").bold = True
        p.add_run(evaluation.next_step)

    if evaluation.flags:
        doc.add_heading("Flags", level=2)
        for flag in evaluation.flags:
            doc.add_paragraph(flag, style="List Bullet")

    doc.add_heading(f"Recommendation: {evaluation.verdict.upper()}", level=2)
    for guidance in evaluation.framing_guidance:
        doc.add_paragraph(guidance, style="List Bullet")

    if evaluation.cover_letter_strategy:
        doc.add_heading("Cover letter strategy", level=2)
        doc.add_paragraph(evaluation.cover_letter_strategy)

    if evaluation.interview_prep:
        doc.add_heading("Interview prep", level=2)
        for tip in evaluation.interview_prep:
            doc.add_paragraph(tip, style="List Bullet")

    out_path = (
        _job_folder(out_dir, company=company, title=title, multi_lead=multi_lead, sibling_titles=sibling_titles)
        / "full-LLM-review.docx"
    )
    doc.save(str(out_path))
    return out_path


def _generate_content(
    jd_text: str,
    *,
    company: str,
    title: str,
    model: str,
    client=None,
) -> tuple[dict, list[CallMetrics]]:
    profile = _load_candidate_profile()
    system = _GENERATE_SYSTEM_PROMPT.replace("{profile}", profile)
    user = f"Company: {company}\nTitle: {title}\n\nJob description:\n{jd_text}"
    return _call_and_parse_json(system, user, model=model, client=client, max_tokens=8192, step="generate")


_HOUSE_RULE_REPAIR_SYSTEM_PROMPT = """Rewrite the JSON object below, fixing ONLY these violations, and leave \
everything else exactly as-is:
{issues}

Specifically: remove any sentence or clause about citizenship, residency, green card, work authorization, or \
security-clearance eligibility (do not replace it with anything, just remove it and smooth the surrounding \
sentence); and remove any dollar figure, hourly rate, salary range, or compensation comparison (e.g. a \
W2-vs-C2C rate comparison) — again just remove it and smooth the surrounding sentence, don't replace it \
with a vaguer restatement of the same idea.

Respond with ONLY the corrected raw JSON object, same schema and keys, no markdown fences, no prose."""


def _repair_house_rule_violations(
    content: dict, *, issues: list[str], model: str, client=None
) -> tuple[dict, list[CallMetrics]]:
    system = _HOUSE_RULE_REPAIR_SYSTEM_PROMPT.format(issues="\n".join(f"- {i}" for i in issues))
    try:
        return _call_and_parse_json(
            system, json.dumps(content), model=model, client=client, max_tokens=8192, step="generate_house_rule_repair"
        )
    except (json.JSONDecodeError, LLMApplyError) as exc:
        logger.warning("House-rule repair pass failed (%s); keeping original content with warnings intact", exc)
        return content, []


def _check_house_rules_text(text: str) -> list[str]:
    """Same mechanical safety net as `_check_house_rules`, for callers that
    generate plain text rather than the résumé/cover-letter JSON shape
    (see `generate_followup_message` below)."""
    warnings = [f"banned term found: {term!r}" for term in _BANNED_TERMS if term.lower() in text.lower()]
    if _WORK_AUTH_RE.search(text):
        warnings.append(
            "possible work-authorization/citizenship/clearance statement found — remove before sending"
        )
    if _COMP_FIGURE_RE.search(text):
        warnings.append("possible compensation figure/rate/range found — remove before sending")
    return warnings


def _check_house_rules(content: dict, *, company: str) -> list[str]:
    """Mechanical safety net for the house rules the generation prompt states
    in natural language. Returns human-readable warning strings; never
    raises, and never silently edits content — callers should surface these
    to the user so they can review before sending anything out."""
    text = json.dumps(content)
    warnings = [f"banned term found: {term!r}" for term in _BANNED_TERMS if term.lower() in text.lower()]

    cover_text = json.dumps(content.get("cover_letter") or {})
    if _WORK_AUTH_RE.search(cover_text):
        warnings.append(
            "possible work-authorization/citizenship/clearance statement found in cover letter — "
            "CLAUDE.md bans these outright; remove before sending"
        )

    if _COMP_FIGURE_RE.search(text):
        warnings.append(
            "possible compensation figure/rate/range found — CLAUDE.md §4 rule 12 bans these outright "
            "(compensation is a live conversation, never written into the package); remove before sending"
        )

    return warnings


_FOLLOWUP_KINDS = ("thank_you", "status_check_in")

_THANK_YOU_SYSTEM_PROMPT = """You are drafting a short, warm, professional thank-you email from {name} to a \
contact after a job interview. 3-5 sentences: thank them for their time, reference any specific discussion \
points given below if there are any, reaffirm genuine interest in the role, and note you're glad to answer any \
follow-up questions. No subject line. Start with "Hi <FirstName>," (or "Hi," if no name is given) and sign off \
with "{name}" on its own line at the end — no other placeholder brackets anywhere in the output. Never invent \
interview details beyond what's given below. Never state a work-authorization/citizenship/clearance position, \
and never mention a dollar figure, hourly rate, salary, or compensation range.

Candidate background (for tone/context only — do not summarize or dump this into the letter):
{profile}"""

_STATUS_CHECK_IN_SYSTEM_PROMPT = """You are drafting a brief, polite check-in email from {name} to a recruiter \
or hiring contact, following up because some time has passed with no update on an application. 2-4 sentences: \
reference the role by name, note it's been a while since the last update, ask for a status update, and \
reaffirm continued interest — no pressure, impatience, or ultimatum in tone. No subject line. Start with "Hi \
<FirstName>," (or "Hi," if no name is given) and sign off with "{name}" on its own line at the end — no other \
placeholder brackets anywhere in the output. Never state a work-authorization/citizenship/clearance position, \
and never mention a dollar figure, hourly rate, salary, or compensation range.

Candidate background (for tone/context only — do not summarize or dump this into the letter):
{profile}"""


@dataclass
class FollowupMessageResult:
    kind: str  # "thank_you" | "status_check_in"
    text: str
    warnings: list[str] = field(default_factory=list)
    metrics: CallMetrics | None = None


def generate_followup_message(
    kind: str,
    *,
    company: str,
    title: str,
    contact_name: str = "",
    context: str = "",
    days_since_contact: int | None = None,
    model: str = DEFAULT_MODEL,
    client=None,
) -> FollowupMessageResult:
    """Draft a short follow-up email (never sent automatically — the caller
    saves/prints it for you to review and send by hand, same human-in-the-
    loop principle as the withdrawal-note idea in docs/JOB_CRM_VISION.md
    UC-8). `kind`:
      - "thank_you": post-interview thank-you note.
      - "status_check_in": a "haven't heard back in a while" nudge, meant
        to pair with `store.find_recent_rejection`'s sibling concept —
        `awaiting_response_since`/UC-6 — for deciding *when* one's overdue.
    """
    if kind not in _FOLLOWUP_KINDS:
        raise ValueError(f"kind must be one of {_FOLLOWUP_KINDS}, got {kind!r}")

    profile = _load_candidate_profile()
    template = _THANK_YOU_SYSTEM_PROMPT if kind == "thank_you" else _STATUS_CHECK_IN_SYSTEM_PROMPT
    system = template.format(name=CANDIDATE_NAME, profile=profile)

    user_lines = [f"Company: {company}", f"Role: {title}"]
    if contact_name:
        user_lines.append(f"Contact name: {contact_name}")
    if days_since_contact is not None:
        user_lines.append(f"Days since last contact: {days_since_contact}")
    if context:
        user_lines.append(f"Additional context: {context}")

    text, metrics = _call(
        system, "\n".join(user_lines), model=model, client=client, max_tokens=1024, step="generate_followup"
    )
    return FollowupMessageResult(kind=kind, text=text.strip(), warnings=_check_house_rules_text(text), metrics=metrics)


def _safe_filename(name: str) -> str:
    return _UNSAFE_FILENAME_CHARS.sub("", name.replace(" ", "_"))


def _job_folder(
    out_dir: Path,
    *,
    company: str,
    title: str,
    multi_lead: bool = False,
    sibling_titles: tuple[str, ...] = (),
) -> Path:
    """Where this job's artifacts (JD, LLM review, résumé, cover letter)
    live together, instead of scattering them across parallel Reviews/
    CoverLetters folders. Layout (2026-07-11 restructure, one `<Company>/`
    folder per company under the output root):
    - `multi_lead=False` (this is the only lead tracked for this company):
      files land flat, directly in `<Company>/`.
    - `multi_lead=True` (this company has 2+ tracked leads): files land in
      `<Company>/<Company>_<Title>/`, one subfolder per lead, so different
      roles at the same company never collide.

    `multi_lead`/`sibling_titles` are computed by the caller from the leads
    DB (see store.get_sibling_titles) since this module has no DB
    connection of its own; default `multi_lead=False` keeps this usable
    without a DB (e.g. tests, one-off scripts) at the cost of not knowing
    about sibling leads.

    Created if missing; reused as-is if a prior evaluate/generate run
    already made it (e.g. the JD + review land first, the résumé/cover
    letter land later only if the verdict is "pursue"). If a second lead
    just appeared for a company that was previously flat (single-lead),
    and there's exactly one unambiguous sibling to attribute the existing
    flat files to, they're migrated into their own subfolder first so
    nothing from the two roles gets mixed together; a less-clear case (0 or
    2+ apparent siblings) is left alone with a warning rather than guessing
    wrong and silently mixing two roles' files."""
    company_dir = out_dir / _safe_filename(company)
    if not multi_lead:
        company_dir.mkdir(parents=True, exist_ok=True)
        return company_dir

    lead_dir = company_dir / _safe_filename(f"{company}_{title}")
    if not lead_dir.exists() and company_dir.exists():
        flat_files = [p for p in company_dir.iterdir() if p.is_file()]
        has_subfolders = any(p.is_dir() for p in company_dir.iterdir())
        if flat_files and not has_subfolders:
            if len(sibling_titles) == 1:
                old_dir = company_dir / _safe_filename(f"{company}_{sibling_titles[0]}")
                old_dir.mkdir(parents=True, exist_ok=True)
                for f in flat_files:
                    f.rename(old_dir / f.name)
                logger.info(
                    "Migrated %s's existing flat files into %s now that a second lead exists",
                    company, old_dir,
                )
            else:
                logger.warning(
                    "%s has flat lead files directly in %s but %d sibling lead(s) in the DB — "
                    "couldn't safely auto-migrate (ambiguous which lead they belong to); "
                    "move them into their own subfolder manually.",
                    company, company_dir, len(sibling_titles),
                )
    lead_dir.mkdir(parents=True, exist_ok=True)
    return lead_dir


def render_job_description(
    jd_text: str,
    *,
    company: str,
    title: str,
    out_dir: Path = DEFAULT_OUTPUT_ROOT,
    multi_lead: bool = False,
    sibling_titles: tuple[str, ...] = (),
) -> Path:
    """Save the JD text this lead was actually evaluated against as its own
    docx, alongside the review/résumé/cover letter — so the full context for
    a decision is preserved even if the source email is later deleted or the
    ATS posting expires."""
    doc = Document()
    doc.add_heading(f"{title} @ {company}", level=1)
    for para in (jd_text or "(no JD text captured)").split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    out_path = (
        _job_folder(out_dir, company=company, title=title, multi_lead=multi_lead, sibling_titles=sibling_titles)
        / "JobDescription.docx"
    )
    doc.save(str(out_path))
    return out_path


def render_resume(
    resume: dict,
    *,
    company: str,
    title: str,
    out_dir: Path = DEFAULT_OUTPUT_ROOT,
    multi_lead: bool = False,
    sibling_titles: tuple[str, ...] = (),
) -> Path:
    doc = Document()
    apply_template_styles(doc)
    doc.add_heading(CANDIDATE_NAME, level=1)
    doc.add_paragraph(resume.get("positioning_line") or "Senior Software Engineer & Independent Consultant")
    # Résumé header omits phone per CLAUDE.md §1 — email/LinkedIn/GitHub only.
    add_muted_contact_line(doc, [CANDIDATE_EMAIL, CANDIDATE_LINKEDIN, CANDIDATE_GITHUB])

    if resume.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(resume["summary"])

    if resume.get("skills"):
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(resume["skills"]))

    experience = list(resume.get("experience") or [])
    # Safety net: HomePortfolio must be last regardless of what the model returned.
    experience.sort(key=lambda e: 1 if "homeportfolio" in str(e.get("employer", "")).lower() else 0)

    if experience:
        doc.add_heading("Experience", level=2)
    for entry in experience:
        header_p = doc.add_paragraph()
        header_p.add_run(str(entry.get("employer", ""))).bold = True
        if entry.get("dates"):
            header_p.add_run(f"   ({entry['dates']})").italic = True
        if entry.get("role_note"):
            note_p = doc.add_paragraph(str(entry["role_note"]))
            for run in note_p.runs:
                run.italic = True
        for bullet in entry.get("bullets") or []:
            doc.add_paragraph(str(bullet), style="List Bullet")
        for sub in entry.get("subsections") or []:
            sub_p = doc.add_paragraph()
            sub_p.add_run(str(sub.get("heading", ""))).bold = True
            for bullet in sub.get("bullets") or []:
                doc.add_paragraph(str(bullet), style="List Bullet")

    if resume.get("education"):
        doc.add_heading("Education & Credentials", level=2)
        for item in resume["education"]:
            doc.add_paragraph(str(item), style="List Bullet")

    candidate = _safe_filename(CANDIDATE_NAME)
    out_path = _job_folder(
        out_dir, company=company, title=title, multi_lead=multi_lead, sibling_titles=sibling_titles
    ) / _safe_filename(f"{candidate}_Resume_{company}_{title}.docx")
    doc.save(str(out_path))
    return out_path


def render_cover_letter(
    cover_letter: dict,
    *,
    company: str,
    title: str,
    out_dir: Path = DEFAULT_OUTPUT_ROOT,
    multi_lead: bool = False,
    sibling_titles: tuple[str, ...] = (),
) -> Path:
    doc = Document()
    apply_template_styles(doc)
    doc.add_heading(CANDIDATE_NAME, level=1)
    add_muted_contact_line(doc, [CANDIDATE_EMAIL, CANDIDATE_PHONE, CANDIDATE_LINKEDIN, CANDIDATE_GITHUB])

    doc.add_paragraph(cover_letter.get("salutation") or "Dear Hiring Team,")
    for para in cover_letter.get("paragraphs") or []:
        text = str(para)
        if looks_like_signature_line(text, phone=CANDIDATE_PHONE, email=CANDIDATE_EMAIL):
            # The model occasionally emits a redundant name/contact line as
            # the first body paragraph, duplicating the header block above
            # (observed bug, 2026-07-12) — drop it rather than render it.
            logger.warning("Dropping duplicate contact/signature line from cover letter body: %r", text)
            continue
        doc.add_paragraph(text)

    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(CANDIDATE_NAME)

    candidate = _safe_filename(CANDIDATE_NAME)
    out_path = _job_folder(
        out_dir, company=company, title=title, multi_lead=multi_lead, sibling_titles=sibling_titles
    ) / _safe_filename(f"{candidate}_coverLetter_{company}_{title}.docx")
    doc.save(str(out_path))
    return out_path


def generate_package(
    jd_text: str,
    *,
    company: str,
    title: str,
    model: str = DEFAULT_MODEL,
    client=None,
    output_root: Path = DEFAULT_OUTPUT_ROOT,
    force: bool = False,
    multi_lead: bool = False,
    sibling_titles: tuple[str, ...] = (),
) -> PackageResult:
    """Evaluate a lead (always saving the JD text + LLM review into that
    job's folder — see `_job_folder`) and, on a "pursue" verdict (or any
    verdict when `force=True`), additionally generate + save a tailored
    résumé and cover letter into the same folder. Returns the evaluation
    either way — CLAUDE.md's "on a dealbreaker or pass, report the mismatch
    and stop" only means "don't generate a résumé/cover letter", not "don't
    write down the review".

    `force` exists for the human-already-decided case (e.g. a "review"
    verdict driven by missing/unwritten JD specifics rather than an actual
    dealbreaker, and the candidate already told the recruiter/hiring contact
    they'd send a résumé) — it does NOT bypass `_check_house_rules`, so
    banned content (comp figures, work-auth statements, etc.) is still
    caught and repaired the same as any other generated package.

    `multi_lead`/`sibling_titles` decide the on-disk folder layout — see
    `_job_folder`. This function has no DB connection of its own, so the
    caller computes these from the leads DB (store.get_sibling_titles) and
    passes them through; the defaults (flat, no known siblings) are only
    right for a standalone/no-DB caller.
    """
    evaluation = evaluate_lead(jd_text, company=company, title=title, model=model, client=client)
    jd_path = render_job_description(
        jd_text, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )
    review_path = render_jd_review_docx(
        evaluation, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )
    if evaluation.verdict != "pursue" and not force:
        return PackageResult(evaluation=evaluation, jd_path=jd_path, review_path=review_path)

    content, generate_calls = _generate_content(jd_text, company=company, title=title, model=model, client=client)
    warnings = _check_house_rules(content, company=company)
    if warnings:
        logger.warning("Generated content for %s / %s failed house-rule checks: %s", company, title, warnings)
        content, repair_calls = _repair_house_rule_violations(content, issues=warnings, model=model, client=client)
        generate_calls += repair_calls
        warnings = _check_house_rules(content, company=company)
        if warnings:
            logger.warning("House-rule violations persisted after repair pass for %s / %s: %s", company, title, warnings)

    resume_path = render_resume(
        content.get("resume") or {}, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )
    cover_letter_path = render_cover_letter(
        content.get("cover_letter") or {}, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )
    return PackageResult(
        evaluation=evaluation,
        jd_path=jd_path,
        review_path=review_path,
        resume_path=resume_path,
        cover_letter_path=cover_letter_path,
        warnings=warnings,
        generate_metrics=_sum_metrics("generate", model, generate_calls),
    )


@dataclass
class TwoTierPackageResult:
    """Result of the two-tier review pipeline (2026-07-11): a free
    rule-based pass always runs first (`no_llm_score`/`no_llm_review_path`,
    saved as `no-LLM-review.docx`); only once that clears
    `scoring.scorer.should_run_llm_review` (or `force=True`) does an
    LLM-backed second pass run (`evaluation`/`full_llm_review_path`, saved
    as `full-LLM-review.docx`); only once *that* comes back "pursue" (or
    `force=True`) does a résumé/cover-letter package get generated."""

    no_llm_score: ScoreResult
    jd_path: Path | None = None
    no_llm_review_path: Path | None = None
    ran_full_llm_review: bool = False
    evaluation: EvaluationResult | None = None
    full_llm_review_path: Path | None = None
    resume_path: Path | None = None
    cover_letter_path: Path | None = None
    warnings: list[str] = field(default_factory=list)
    generate_metrics: CallMetrics | None = None

    @property
    def total_input_tokens(self) -> int:
        return (self.evaluation.metrics.input_tokens if self.evaluation and self.evaluation.metrics else 0) + (
            self.generate_metrics.input_tokens if self.generate_metrics else 0
        )

    @property
    def total_output_tokens(self) -> int:
        return (self.evaluation.metrics.output_tokens if self.evaluation and self.evaluation.metrics else 0) + (
            self.generate_metrics.output_tokens if self.generate_metrics else 0
        )

    @property
    def total_elapsed_s(self) -> float:
        return (self.evaluation.metrics.elapsed_s if self.evaluation and self.evaluation.metrics else 0.0) + (
            self.generate_metrics.elapsed_s if self.generate_metrics else 0.0
        )

    @property
    def total_cost_usd(self) -> float | None:
        parts = [
            m.cost_usd
            for m in ((self.evaluation.metrics if self.evaluation else None), self.generate_metrics)
            if m is not None and m.cost_usd is not None
        ]
        return sum(parts) if parts else None


def generate_two_tier_package(
    jd_text: str,
    *,
    company: str,
    title: str,
    model: str = DEFAULT_MODEL,
    client=None,
    output_root: Path = DEFAULT_OUTPUT_ROOT,
    force: bool = False,
    force_llm_review: bool = False,
    generate: bool = True,
    multi_lead: bool = False,
    sibling_titles: tuple[str, ...] = (),
) -> TwoTierPackageResult:
    """Two-tier review pipeline (2026-07-11), replacing `generate_package`'s
    always-LLM flow for real pipeline use:

    1. Always run the free, deterministic rule-based scorer
       (`scoring.scorer.score_jd`) and save `no-LLM-review.docx`.
    2. Only if that score clears `scoring.scorer.should_run_llm_review`
       (or `force=True`/`force_llm_review=True`), spend an LLM call:
       `evaluate_lead` + save `full-LLM-review.docx`.
    3. Only if the LLM verdict is "pursue" (or `force=True`) AND
       `generate=True`, generate + save a tailored résumé and cover letter.

    `force=True` bypasses both gates — the human already decided this
    specific lead is worth the full treatment regardless of what the free
    pass says (e.g. a manual `apply_package.py --force` run on one lead
    someone already wants a full review of). It does NOT bypass
    `_check_house_rules` — banned content is still caught and repaired the
    same as any other generated package.

    `force_llm_review=True` bypasses only gate 2 (still respects gate 3 —
    résumé/cover letter only get generated on an actual "pursue" verdict
    from the LLM). Useful when the caller wants the LLM's nuanced judgment
    regardless of the free pass's score, but doesn't want to blindly
    generate application docs the way a bare `force=True` would.

    `generate=False` (e.g. `triage_recruiter_inbox.py --no-generate`)
    still respects both gates above for scoring purposes — it only
    additionally skips step 3's résumé/cover-letter generation, so a
    "score with the LLM but never spend on generation" preview still costs
    an LLM call only when the free pass says it's worth it.
    """
    jd_path = render_job_description(
        jd_text, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )
    no_llm_score = score_jd(jd_text)
    no_llm_review_path = render_no_llm_review_docx(
        no_llm_score, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )
    result = TwoTierPackageResult(no_llm_score=no_llm_score, jd_path=jd_path, no_llm_review_path=no_llm_review_path)

    if not (should_run_llm_review(no_llm_score) or force or force_llm_review):
        return result

    result.ran_full_llm_review = True
    evaluation = evaluate_lead(jd_text, company=company, title=title, model=model, client=client)
    result.evaluation = evaluation
    result.full_llm_review_path = render_jd_review_docx(
        evaluation, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )

    if (evaluation.verdict != "pursue" and not force) or not generate:
        return result

    content, generate_calls = _generate_content(jd_text, company=company, title=title, model=model, client=client)
    warnings = _check_house_rules(content, company=company)
    if warnings:
        logger.warning("Generated content for %s / %s failed house-rule checks: %s", company, title, warnings)
        content, repair_calls = _repair_house_rule_violations(content, issues=warnings, model=model, client=client)
        generate_calls += repair_calls
        warnings = _check_house_rules(content, company=company)
        if warnings:
            logger.warning("House-rule violations persisted after repair pass for %s / %s: %s", company, title, warnings)

    result.resume_path = render_resume(
        content.get("resume") or {}, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )
    result.cover_letter_path = render_cover_letter(
        content.get("cover_letter") or {}, company=company, title=title, out_dir=output_root,
        multi_lead=multi_lead, sibling_titles=sibling_titles,
    )
    result.warnings = warnings
    result.generate_metrics = _sum_metrics("generate", model, generate_calls)
    return result
