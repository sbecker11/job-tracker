"""Shared style constants + helpers for résumé/cover-letter `.docx` generation.

Single source of truth: `~/Desktop/Resumes/2026/Templates/Shawn_Becker_Template.dotx`.
That template's stylesheet is the full standard Word style catalog (Heading1-9,
ListBullet, BodyText, etc.) with only `Heading1` actually customized — the same
name/contact/body treatment applies equally well to résumés and cover letters,
so both document types share this one template/module rather than each having
their own.

python-docx cannot open a `.dotx` as a base document (it rejects the
`...wordprocessingml.template.main+xml` content type outright), so instead of
binding to that file at runtime, the values below were extracted from its
`word/styles.xml` / `document.xml` (via `unzip`, 2026-07-12) and are re-applied
through python-docx's API onto the `Normal` / `Heading 1` styles and section
page setup of every freshly-built `Document()`. If the template's look ever
changes, update the constants here — nowhere else in the generator should
hardcode a font, size, or color for these two document types.
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Arial"
BODY_SIZE = Pt(11)

HEADING_FONT = "Arial"
HEADING_SIZE = Pt(18)
HEADING_COLOR_HEX = "365F91"

MUTED_FONT = "Arial"
MUTED_SIZE = Pt(10)
MUTED_COLOR_HEX = "555555"

# Template docDefaults: spacing after=200 (dxa, i.e. 10pt) / line=276 lineRule="auto"
# (276/240 = 1.15x). Paragraph spacing comes from these style defaults, not
# from empty spacer paragraphs — see `render_cover_letter`/`render_resume`.
PARAGRAPH_SPACE_AFTER = Pt(10)
PARAGRAPH_LINE_SPACING = 1.15

# Template sectPr: pgMar top=1440 bottom=1440 left=1800 right=1800 header=720
# footer=720 (dxa; 1440 dxa = 1"). Page size (12240x15840 dxa = US Letter) is
# already python-docx's default, so it's not set explicitly here.
PAGE_MARGIN_TOP = Inches(1)
PAGE_MARGIN_BOTTOM = Inches(1)
PAGE_MARGIN_LEFT = Inches(1.25)
PAGE_MARGIN_RIGHT = Inches(1.25)
PAGE_MARGIN_HEADER = Inches(0.5)
PAGE_MARGIN_FOOTER = Inches(0.5)


def apply_template_styles(doc: Document) -> None:
    """Apply the canonical Shawn_Becker_Template.dotx look to a freshly
    created python-docx `Document()`: body font/spacing on `Normal`, the
    Heading-1 name treatment, and page margins. Call this once per document,
    immediately after `Document()` and before adding any content, so every
    paragraph/heading added afterward inherits it via style rather than
    needing direct per-run formatting."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = PARAGRAPH_SPACE_AFTER
    normal.paragraph_format.line_spacing = PARAGRAPH_LINE_SPACING
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = HEADING_FONT
    heading1.font.size = HEADING_SIZE
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor.from_string(HEADING_COLOR_HEX)
    heading1.paragraph_format.space_after = Pt(0)

    section = doc.sections[0]
    section.top_margin = PAGE_MARGIN_TOP
    section.bottom_margin = PAGE_MARGIN_BOTTOM
    section.left_margin = PAGE_MARGIN_LEFT
    section.right_margin = PAGE_MARGIN_RIGHT
    section.header_distance = PAGE_MARGIN_HEADER
    section.footer_distance = PAGE_MARGIN_FOOTER


def add_muted_contact_line(doc: Document, segments: list[str]) -> None:
    """Add the standard muted Arial 10pt #555555 contact line, joining
    `segments` with '  |  ' — every segment gets the muted treatment (fixes
    the historical bug where only the last/GitHub segment was styled and
    the rest fell through to default black 11pt)."""
    p = doc.add_paragraph()
    run = p.add_run("  |  ".join(segments))
    run.font.name = MUTED_FONT
    run.font.size = MUTED_SIZE
    run.font.color.rgb = RGBColor.from_string(MUTED_COLOR_HEX)


def looks_like_signature_line(text: str, *, phone: str, email: str) -> bool:
    """Heuristic: does this look like a redundant name/contact signature
    line the model sometimes generates as a stray first body paragraph
    (duplicating the header block)? Matching on the literal phone number or
    email address is deliberately narrow — both are unique, unlikely to
    appear in genuine cover-letter prose, so this has essentially no
    false-positive risk while reliably catching the duplicate-header bug."""
    t = text.strip()
    return bool(t) and (phone in t or email in t)
